import pyautogui;
import datetime
from docx import Document
import time
import os
import glob


current_directory = os.getcwd()
final_directory = os.path.join(current_directory, r'images')
if not os.path.exists(final_directory):
   os.makedirs(final_directory)

path = final_directory+"\\*.png";


print(datetime.datetime.now())

Document().save('screenshot.docx')

while(1):
    inn=input('enter to take new screenshot, "enter create to create document"');
    if(len(inn)>0):
        break
    pyautogui.screenshot()
    pyautogui.screenshot(final_directory+'\\'+str(int(time.time()))+'.png')


document = Document('screenshot.docx')


for fname in glob.glob(path):
    print(fname);
    p = document.add_paragraph()
    r = p.add_run()
    r.add_picture(fname)

document.save('screenshot.docx')

print(datetime.datetime.now())
